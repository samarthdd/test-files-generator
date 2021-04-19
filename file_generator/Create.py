import xlsxwriter
import ntpath

from osbot_utils.utils.Files import file_create
from fpdf                    import FPDF
from docx                    import Document
from PIL                     import Image, ImageDraw, ImageFont

class Create(object):

    def create(self, path, file_type, content):

        self.file_type = file_type
        self.content   = content
        self.path = path

        if len(ntpath.basename(path).split(".")) == 1:
            self.path = self.path + "." + file_type

        if file_type == "png" or file_type == "jpg" or file_type == "jpeg" :
            method_name = "create_image"
        else:
            method_name = "create_" + self.file_type

        method = getattr(self, method_name, lambda: 'Invalid File Type')
        return method()

    def create_txt(self):
        file_create(path=self.path, contents=self.content)

    def create_pdf(self):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_xy(0, 0)
        pdf.set_font('arial', 'B', 12.0)
        pdf.multi_cell(0, 5, txt=self.content)
        pdf.output(self.path, 'F')
        return

    def create_docx(self):
        document = Document()
        document.add_heading('This is the title', 0)
        document.add_paragraph(self.content)
        document.save(self.path)
        return


    def create_xlsx(self):
        workbook = xlsxwriter.Workbook(self.path)
        worksheet = workbook.add_worksheet()
        worksheet.write('A1', 'Content')
        worksheet.write('B1', self.content)
        workbook.close()
        return

    def create_image(self):
        img = Image.new('RGB', (500, 200), color=(73, 109, 137))
        d = ImageDraw.Draw(img)
        d.text((10, 10),  self.content, fill=(255, 255, 0))
        img.save(self.path)
        return

    def create_gif(self):
        images = []
        width = 10
        center = width // 2
        color_1 = (0, 0, 0)
        color_2 = (255, 255, 255)
        max_radius = int(center * 1.5)
        step = 2

        for i in range(0, max_radius, step):
            im = Image.new('RGB', (width, width), color_1)
            draw = ImageDraw.Draw(im)
            draw.ellipse((center - i, center - i, center + i, center + i), fill=color_2)
            images.append(im)
        for i in range(0, max_radius, step):
            im = Image.new('RGB', (width, width), color_2)
            draw = ImageDraw.Draw(im)
            draw.text((10, 10), self.content, fill=(255, 255, 0))
            draw.ellipse((center - i, center - i, center + i, center + i), fill=color_1)
            images.append(im)

        images[0].save(self.path,save_all=True, append_images=images[1:], optimize=False, duration=1, loop=0)
